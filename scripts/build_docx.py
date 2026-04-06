#!/usr/bin/env python3
"""
TIME-AR-001: Build DOCX tables from processed data.
Generates a supplementary DOCX with key data tables.
"""

import csv
import os
import sys

DATA_DIR = os.path.join(os.path.dirname(__file__), '..', 'data', 'processed')
DOCS_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')
FIG_DIR = os.path.join(os.path.dirname(__file__), '..', 'figures')

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def load_csv(filename):
    path = os.path.join(DATA_DIR, filename)
    if not os.path.exists(path):
        print(f"  Warning: {filename} not found, skipping")
        return []
    with open(path) as f:
        return list(csv.DictReader(f))

def build_docx():
    if not HAS_DOCX:
        print("python-docx not available. Building text summary instead.")
        build_text_summary()
        return

    doc = Document()

    # Title
    doc.add_heading('TIME-AR-001: Temporal Admissible Region Study', 0)
    doc.add_paragraph('Supplementary Data Tables — Generated from Processed Data')
    doc.add_paragraph('')

    # Table 1: Classification Summary
    summary = load_csv('classification_summary.csv')
    if summary:
        doc.add_heading('Table 1: F15 Classification Summary', level=1)
        table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
        headers = ['Classification', 'Count', 'Fraction']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in summary:
            cells = table.add_row().cells
            cells[0].text = row.get('classification', '')
            cells[1].text = row.get('count', '')
            cells[2].text = row.get('fraction', '')
        doc.add_paragraph('')

    # Table 2: F15 Scorecard (organism-level)
    scorecard = load_csv('f15_scorecard.csv')
    if scorecard:
        doc.add_heading('Table 2: Organism-Level F15 Scorecard', level=1)
        table = doc.add_table(rows=1, cols=10, style='Light Grid Accent 1')
        headers = ['Organism', 'Domain', 'D_org', 'D_fam', 'Family', 'T_mid', 'gc_org', 'gc_fam', 'F15a', 'F15b']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in scorecard:
            cells = table.add_row().cells
            cells[0].text = row.get('organism', '')
            cells[1].text = row.get('domain', '')
            cells[2].text = row.get('D_organism', '')
            cells[3].text = row.get('D_family', '')
            cells[4].text = row.get('family_name', '')
            cells[5].text = row.get('T_midpoint', '')
            cells[6].text = row.get('gamma_crit_organism', '')
            cells[7].text = row.get('gamma_crit_family', '')
            cells[8].text = row.get('F15a', '')
            cells[9].text = row.get('F15b', '')
        doc.add_paragraph('')

    # Table 3: Gamma_crit ranked table
    gamma = load_csv('gamma_crit_table.csv')
    if gamma:
        doc.add_heading('Table 3: Gamma_crit Ranked Table', level=1)
        cols = list(gamma[0].keys()) if gamma else []
        table = doc.add_table(rows=1, cols=len(cols), style='Light Grid Accent 1')
        for i, h in enumerate(cols):
            table.rows[0].cells[i].text = h
        for row in gamma:
            cells = table.add_row().cells
            for i, h in enumerate(cols):
                cells[i].text = row.get(h, '')
        doc.add_paragraph('')

    # Table 4: Fisher test result
    fisher = load_csv('fisher_test_result.csv')
    if fisher:
        doc.add_heading('Table 4: Fisher Exact Test Result', level=1)
        cols = list(fisher[0].keys()) if fisher else []
        table = doc.add_table(rows=1, cols=len(cols), style='Light Grid Accent 1')
        for i, h in enumerate(cols):
            table.rows[0].cells[i].text = h
        for row in fisher:
            cells = table.add_row().cells
            for i, h in enumerate(cols):
                cells[i].text = row.get(h, '')
        doc.add_paragraph('')

    # Add figures if they exist
    doc.add_heading('Figures', level=1)
    fig_files = [
        ('fig_T1_exclusion_zone.png', 'Figure 1: Temporal Exclusion Zone (D vs log T)'),
        ('fig_T2_gamma_crit_bars.png', 'Figure 2: Gamma_crit for T3req Families'),
        ('fig_T3_deep_families.png', 'Figure 3: Deepest Gene Family Hierarchies'),
        ('fig_T4_physical_vs_bio.png', 'Figure 4: Physical Fractals vs Biological Hierarchies'),
        ('fig_T5_two_regime_gamma.png', 'Figure 5: Two-Regime Empirical Gamma Structure'),
        ('fig_T6_d_distribution.png', 'Figure 6: Gene Family Depth Distribution'),
        ('fig_T7_sensitivity.png', 'Figure 7: Gamma_crit Sensitivity Analysis'),
        ('fig_T8_classification_pie.png', 'Figure 8: F15 Classification Distribution'),
    ]
    for fname, caption in fig_files:
        fpath = os.path.join(FIG_DIR, fname)
        if os.path.exists(fpath):
            doc.add_paragraph(caption, style='Caption')
            doc.add_picture(fpath, width=Inches(5.5))
            doc.add_paragraph('')
        else:
            doc.add_paragraph(f'{caption} — [figure not yet generated]')

    # Save
    out_path = os.path.join(DOCS_DIR, 'TIME_AR_001_Tables.docx')
    doc.save(out_path)
    print(f"DOCX written to: {out_path}")

def build_text_summary():
    """Fallback: write a text summary if python-docx is not available."""
    out_path = os.path.join(DOCS_DIR, 'TIME_AR_001_Tables.txt')

    with open(out_path, 'w') as f:
        f.write("TIME-AR-001: Temporal Admissible Region Study\n")
        f.write("Supplementary Data Tables\n")
        f.write("=" * 60 + "\n\n")

        # Classification summary
        summary = load_csv('classification_summary.csv')
        if summary:
            f.write("Classification Summary\n")
            f.write("-" * 40 + "\n")
            for row in summary:
                f.write(f"  {row.get('classification', '')}: {row.get('count', '')} ({row.get('fraction', '')})\n")
            f.write("\n")

        # F15 scorecard
        scorecard = load_csv('f15_scorecard.csv')
        if scorecard:
            f.write("Organism-Level F15 Scorecard\n")
            f.write("-" * 40 + "\n")
            for row in scorecard:
                f.write(f"  {row.get('organism', '')}: F15a={row.get('F15a', '')} F15b={row.get('F15b', '')}\n")
            f.write("\n")

        # Fisher test
        fisher = load_csv('fisher_test_result.csv')
        if fisher:
            f.write("Fisher Exact Test\n")
            f.write("-" * 40 + "\n")
            for row in fisher:
                f.write(f"  p-value: {row.get('p_value', '')}\n")
                f.write(f"  {row.get('conclusion', '')}\n")
            f.write("\n")

    print(f"Text summary written to: {out_path}")
    print("Note: Install python-docx for full DOCX output: pip3 install python-docx")

if __name__ == '__main__':
    build_docx()
